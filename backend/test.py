from docx import Document

def create_dock(xx_list, zj_list, plan_list):
    # 创建一个新的文档
    doc = Document()

    # 添加标题
    doc.add_heading(text='xxx年xxx月销售分析报告', level=0)

    # 添加标题
    doc.add_heading(text='一：详细分析。', level=1)

    for i in xx_list:
        # 添加列表
        doc.add_paragraph(i, style='ListBullet')

    # 添加标题
    doc.add_heading(text='二：总结分析。', level=1)

    for i in zj_list:
        # 添加列表
        doc.add_paragraph(i, style='ListBullet')

    # 添加标题
    doc.add_heading(text='三：未来销售计划。', level=1)

    for i in plan_list:
        # 添加列表
        doc.add_paragraph(i, style='ListBullet')

    # 保存文档
    doc.save('example.docx')
    print("文档已保存为 'example.docx'")
if __name__ == '__main__':
    xx_list=["1","2","3"]
    zj_list=["4","5","6"]
    plan_list=["7","8","9"]
    create_dock(xx_list, zj_list, plan_list)