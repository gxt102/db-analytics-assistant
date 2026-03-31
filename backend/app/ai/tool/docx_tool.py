from docx import Document
from pydantic import BaseModel,Field
from langchain.tools import tool
import time
import os

#定义一个写入报告的参数类
class DocxArgs(BaseModel):
    title:str = Field(...,description="报告标题")
    xx_list : list = Field(...,description="详细分析")
    zj_list : list = Field(...,description="总结分析")
    plan_list : list = Field(...,description="未来销售计划")

@tool(args_schema=DocxArgs)
def create_dock(title:str,xx_list:list,zj_list:list,plan_list:list) -> str:
    """
    生成word文档
    """
    # 创建一个新的文档
    doc = Document()
    # 添加标题
    doc.add_heading(title, 0)
    # 添加标题
    doc.add_heading('一 ：详细分析。', 1)
    for i in  xx_list:
      # 添加列表
       doc.add_paragraph(i, style='ListBullet')
    # 添加标题
    doc.add_heading('二 ：总结分析。', 1)
    for i in  zj_list:
      # 添加列表
       doc.add_paragraph(i, style='ListBullet')
    # 添加标题
    doc.add_heading('三 ：未来销售计划。', 1)
    for i in plan_list:
        # 添加列表
        doc.add_paragraph(i, style='ListBullet')
    #用当前时间命名文件
    file_name =time.strftime("%Y%m%d%H%M%S", time.localtime())
    #定义文件保存路径,web服务器路径写法
    file_path =os.getcwd()+f"\\static\\download\\{file_name}.docx"
    #main 方法测试
    #file_path = os.getcwd() + f"\\{file_name}.docx"
    # 保存文档
    doc.save(file_path)
    #返回文件下载路径
    down_path =f"http://localhost:8000/static/download/{file_name}.docx"
    return down_path